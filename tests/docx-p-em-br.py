#! python3
from tempfile import NamedTemporaryFile as tempfile
import unittest

from docx import Document  # type: ignore

from topdf import options
from topdf import html_conv_docx as tgt


class TestRunner(unittest.TestCase):
    def test_em(self):
        ftmp1 = tempfile("wt", delete=False)
        ftmp1.write("""
        <html><body><p>abc
        <em>cde</em>
        fgh
        <em>ijk</em> <br />
        lmn
        </p><p><em>a</em>b</p></body></html>
        """)
        fname = ftmp1.name
        ftmp1.close()

        fdest = tempfile("wt", delete=False)
        fdest.close()
        fnam2 = fdest.name

        opts = options.parse([fname, "-o", fnam2])
        tgt.main(opts)

        doc = Document(fnam2)
        output = doc.paragraphs[0].text
        self.assertEqual(output, "abc cde fgh ijk \n lmn ")

        output = doc.paragraphs[1].text
        self.assertEqual(output, "ab")
        self.assertEqual(len(doc.paragraphs), 2)

    def test_strong(self):
        ftmp1 = tempfile("wt", delete=False)
        ftmp1.write("""
        <html><body><p><strong>a</strong>b<br />c
        <strong>d</strong>e<br />
        f<br />g
        </p></body></html>
        """)
        fname = ftmp1.name
        ftmp1.close()

        fdest = tempfile("wt", delete=False)
        fdest.close()
        fnam2 = fdest.name

        opts = options.parse([fname, "-o", fnam2])
        tgt.main(opts)

        doc = Document(fnam2)
        output = doc.paragraphs[0].text
        output = output.replace(" ", "")
        output = output.replace("\n", "")

        self.assertEqual(output, "abcdefg")
        self.assertEqual(len(doc.paragraphs), 1)
