#! python3
import os
from tempfile import NamedTemporaryFile as tempfile
import unittest

from docx import Document  # type: ignore

from topdf import options
from topdf import html_conv_docx as tgt


class Tests(unittest.TestCase):
    def svg_ext(self):
        ftmp1 = tempfile("wt", delete=False)
        ftmp1.write("""
        <html><body><p>abc
        <img src="%s/docx-svg-sample1.svg" />
        cde</p></body></html>
        """ % os.path.dirname(__file__))
        fname = ftmp1.name
        ftmp1.close()

        fnam2 = os.path.basename(__file__) + ".docx"
        opts = options.parse([fname, "-o", fnam2])
        tgt.main(opts)

        doc = Document(fnam2)
        output = doc.paragraphs[0].text
        self.assertEqual(output, "abc cde")

    def svg_in(self):
        ftmp1 = tempfile("wt", delete=False)
        ftmp1.write("""
        <html><body><p>abc
        <svg height="100" width="100">
            <circle cx="50" cy="50" r="40" stroke="black" stroke-width="3"
                fill="red" />
        </svg>
        cde</p></body></html>
        """)
        fname = ftmp1.name
        ftmp1.close()

        fnam2 = os.path.basename(__file__) + ".docx"
        opts = options.parse([fname, "-o", fnam2])
        tgt.main(opts)

        doc = Document(fnam2)
        output = doc.paragraphs[0].text
        self.assertEqual(output, "abc  cde")
