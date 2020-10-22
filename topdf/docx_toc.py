# Copyright (c) 2020, Shimoda <kuri65536 at hotmail dot com>
#
# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at http://mozilla.org/MPL/2.0/.
#
# include {{{1
import re
from typing import Text

from bs4.element import Tag  # type: ignore

from docx import Document  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

try:
    from . import common
except ImportError:
    import common  # type: ignore


def generate_toc(doc: Document, elem: Tag) -> Paragraph:  # {{{1
    def cb(para: Paragraph) -> Paragraph:
        return cache_toc(doc, elem)
        return para

    # [@P5-2-11] a TOC field
    para = doc.add_paragraph()
    common.docx_add_field(para, r'TOC \o "1-9" \h', cb, True)
    return doc.add_paragraph()


def cache_toc(doc: Document, elem: Tag) -> Paragraph:  # {{{1
    para = doc.add_paragraph()
    seq = [i for i in elem.parents if i.name == "body"]
    if len(seq) < 1:
        return para
    # [@P5-2-12] a TOC list
    body = seq[0]
    for elem in body.find_all(re.compile("h[0-9]")):
        level = elem.name.replace("h", "")
        id_ = elem.attrs.get("id", "")
        if len(id_) < 1:
            para.text = elem.text
        else:
            common.docx_add_hlink(para, Text(elem.text), "ahref_" + id_)
        # [@P5-1-12] line spacing of TOC
        para.style = common.docx_style(doc, "TOC Contents " + level)
        para = doc.add_paragraph()
    return para
