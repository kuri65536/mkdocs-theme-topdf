# Copyright (c) 2020, Shimoda <kuri65536 at hotmail dot com>
#
# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at http://mozilla.org/MPL/2.0/.
#
import base64
from logging import (debug as debg, error as eror, warning as warn, )
import os
import sys
from tempfile import NamedTemporaryFile as Temporary
from typing import (Callable, Dict, Iterable, List, Optional, Set,
                    Text, Tuple, )
import zlib

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.text import (                  # type: ignore
        WD_COLOR_INDEX, WD_LINE_SPACING,  # type: ignore
        WD_PARAGRAPH_ALIGNMENT, )             # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Mm, Pt, RGBColor  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from bs4.element import Tag  # type: ignore
import get_image_size  # type: ignore


try:
    import requests
    f_not_online = False
except ImportError:
    f_not_online = True


if False:
    List


class glb:  # {{{1
    numbers_of_captions: Dict[Text, int] = {}
    styles: List[Text] = []
    seq_styles: Dict[Text, Callable[['Styles', Document], Text]] = {}
    seq_files: List[Text] = []
    bookmark_id = 1


class ParseError(Exception):  # {{{1
    pass


def has_class(tag: Tag, *names: Text) -> bool:  # {{{1
    if tag.name is None:
        return False
    classes: Set[Text] = set(tag.attrs.get("class", []))
    targets = set(names)
    return len(classes & targets) > 0


def has_prev_sibling(target: Tag, *tags: Text) -> bool:  # {{{1
    for elem in target.previous_siblings:
        if elem.name is None:
            continue
        if elem.name in tags:
            return True
        break
    return False


def has_next_table(target: Tag) -> bool:  # {{{1
    for elem in target.next_siblings:
        if elem.name is None:
            continue
        if elem.name == "table":
            return True
        break
    parent = target.parent
    if parent is None:
        return False
    for elem in parent.next_siblings:
        if elem.name is None:
            continue
        if elem.name == "p":
            cls = elem.attrs.get("class", [])
            if "before-dl-table" in cls:
                return True
        break
    return False


def has_width_and_parse(classes: Iterable[Text]  # {{{1
                        ) -> Tuple[Text, List[Mm]]:
    def parse_float_or_0(src: Text) -> float:
        try:
            n = float(src)
        except ValueError:
            return 0.0
        return n

    def parse_base(src: Text) -> float:
        n_max, n_sum = 0, 0.0
        for num in src.split("-"):
            n = parse_float_or_0(num)
            if n <= 0.0:
                if num == "a":
                    continue
                if num.endswith("mm") and parse_float_or_0(num[:-2]) > 0.0:
                    continue
                return float("nan")
            n_max, n_sum = max(n, n_max), n_sum + n
        if n_sum > 10.49999:
            return 160.0 / n_sum
        return float(160 / 10)

    src = ""
    for cls in classes:
        if cls.startswith("table"):
            src = cls
            break
    else:
        return "", []
    ret: List[Mm] = []
    base = parse_base(src[5:])
    if base != base:
        return "", []
    for num in src[5:].split("-"):
        if num == "a":
            ret.append(Mm(0))
            continue
        if num.endswith("mm"):
            n = parse_float_or_0(num[:-2])
            ret.append(Mm(n))
            continue
        n = parse_float_or_0(num)
        ret.append(Mm(n * base))
    return src, ret


def count_tags_around_image(src: Tag):  # {{{1
    "[@P14-1-12] image under a rule"
    i = 0
    for elem in src.children:
        if elem.name is None:
            continue
        i += 1
    return i


def classes_from_prev_sibling(target: Tag) -> Iterable[Text]:  # {{{1
    for elem in target.previous_siblings:
        if elem.name is None:
            continue
        ret: List[Text] = elem.attrs.get("class", [])
        return ret + []
    return []


def table_cellspan(e: Tag, *keys: Text) -> Tuple[int, ...]:  # {{{1
    ret: Tuple[int, ...] = ()
    for key in keys:
        ret += ((int(e.get(key)) - 1) if e.has_attr(key) else 0, )
    return ret


def table_update_rowcolspan(dct: Dict[Tuple[int, int], Tag]  # {{{1
                            ) -> Dict[Tuple[int, int], Tag]:
    """[@P13-1-13] alignment cell potisions"""
    rowspans: List[Set[int]] = [set()]
    ret: Dict[Tuple[int, int], Tag] = {}
    r, c = -1, 0
    for (row, col), elem in sorted(dct.items(), key=lambda x: x[0]):
        if r != row:
            rowspans = rowspans[1:]
            rowspans = rowspans if len(rowspans) > 0 else [set()]
        r, c = row, col
        while c in rowspans[0]:
            c += 1
        x, y = table_cellspan(elem, "colspan", "rowspan")
        cols: Set[int] = {c}
        ret[(r, c)] = elem
        if x != 0:
            for i in range(1, x + 1):
                cols.add(c + i)
        rowspans[0].update(cols)
        if y != 0:
            for i in range(1, y + 1):
                if i >= len(rowspans):
                    rowspans.append(cols)
                else:
                    rowspans[i].update(cols)
    """while True:  # dump cells information
        r, msg = -1, ""
        for (row, col), elem in sorted(ret.items(), key=lambda x: x[0]):
            x, y = table_cellspan(elem, "colspan", "rowspan")
            if r != row:
                warn("common.table:" + msg[1:])
                msg, r = "", row
            msg += ",(%d,%d" % (row, col)
            if x > 0 or y > 0:
                msg += "-%d,%d" % (x, y)
            msg += ")"
        if len(msg):
            warn("common.table:" + msg[1:])
        break"""
    return ret


def is_online_mode() -> bool:  # {{{1
    import socket
    socket.setdefaulttimeout(1.0)
    try:
        socket.gethostbyname("www.python.org")
        return True
    except OSError:
        pass
    return False


def init(force_offline: bool) -> None:  # {{{1
    global f_not_online
    if f_not_online:
        warn("img-tag: package 'requests' required, "
             "online images will be ignored...")
        return

    if force_offline is not True and is_online_mode():
        f_not_online = False
    else:
        warn("img-tag: now in off-line mode, "
             "online images will be ignored...")
        f_not_online = True


def is_target_in_http(url: Text) -> bool:  # {{{1
    if url.startswith("http://"):
        return True
    elif url.startswith("https://"):
        return True
    return False


def remove_temporaries() -> None:  # {{{1
    for fname in glb.seq_files:
        try:
            os.remove(fname)
        except OSError:
            pass
    try:
        os.rmdir("tmp")
    except OSError:
        pass


def download_image(url_doc: Text, src: Text) -> Text:  # {{{1
    # absolute
    if src.startswith("http://"):
        fname = download_image_run(src)
    elif src.startswith("https://"):
        fname = download_image_run(src)
    elif src.startswith("file://"):
        fname = src[7:]  # just remove prefix
    elif src.startswith("data:image/"):
        fname = download_image_extract(src)
    elif "://" in src:
        eror("img-tag: url was not recognized, ignored...: " + src)
        return ""
    else:  # relative
        if is_target_in_http(url_doc):
            src = download_image_unified_http(url_doc, src)
            fname = download_image_run(src)
        else:
            dname = os.path.dirname(url_doc)
            fname = os.path.join(dname, src)
            fname = os.path.realpath(fname)

    if os.path.exists(fname):
        return fname
    warn("img-tag: file is not found: " + fname)
    return ""


def download_image_run(url: Text) -> Text:  # {{{1
    if f_not_online:
        eror("img-tag: now offline, ignored: " + url)
        return ""

    sfx = url.split("?")[0]   # remove the request string.
    sfx = sfx.split(".")[-1]  # remove prefix.
    if sfx.lower() not in ("jpg", "png", "svg", "tiff", "tif", "gif"):
        eror("img-tag: unsupported format: " + sfx)
        return ""
    sfx = "." + sfx

    resp = requests.get(url)
    if (resp.status_code / 100) != 2:
        return ""  # failed to download

    dname, fname = "tmp", ""
    if not os.path.exists(dname):
        os.mkdir(dname)
    with Temporary(mode="w+b", dir=dname, suffix=sfx, delete=False
                   ) as fp:
        fname = fp.name
        fp.write(resp.content)
    glb.seq_files.append(fname)
    return fname


def download_image_unified_http(url: Text, src: Text) -> Text:  # {{{1
    # TODO(shimoda): check target url and manipulate url of images.
    if url.endswith("/"):
        return url + src
    seq = url.split("/")
    seq = seq[:-1]
    seq.append(src)
    return "/".join(seq)


def download_image_extract(url: Text) -> Text:  # {{{1
    url = url[11:]  # remove "data:image/"
    if ";" not in url:
        return ""
    i = url.index(";")
    sfx = "." + url[:i]
    url = url[i:]
    if "," not in url:
        return ""
    i = url.index(",")
    if not url.startswith(";base64,"):
        warn("img-tag: now support extract from base64: " + url[:i])
        return ""
    url = url[i:]
    data = base64.b64decode(url)

    dname, fname = "tmp", ""
    if not os.path.exists(dname):
        os.mkdir(dname)
    with Temporary(mode="w+b", dir=dname, suffix=sfx, delete=False
                   ) as fp:
        fname = fp.name
        fp.write(data)
    glb.seq_files.append(fname)
    return fname


def image_width_and_height(src: Text) -> Tuple[int, int]:  # {{{1
    try:
        return get_image_size.get_image_size(src)  # type: ignore
    except get_image_size.UnknownImageFormat:
        pass
    return (-1, -1)


def dot_to_mm(n: int) -> int:  # {{{1
    inch = n / 96         # dpi -> inch
    mm = Mm(inch * 25.4)  # inch -> Mm
    return mm  # type: ignore


def dot_to_page(w: int, h: int) -> Dict[Text, int]:  # {{{1
    if dot_to_mm(w) > Mm(210 - 40):
        args = {"width": Mm(210 - 40)}
        if dot_to_mm(h) > Mm(279 - 40):
            if h > w:
                args = {"height": Mm(279 - 40)}
    elif dot_to_mm(h) > Mm(279 - 40):
        args = {"height": Mm(279 - 40)}
    else:
        args = {"width": dot_to_mm(w),
                "height": dot_to_mm(h)}
    return args


def docx_add_field(para: Paragraph, instr: Text,  # {{{1
                   cache: Callable[[Paragraph], Paragraph],
                   dirty: Optional[bool] = None) -> None:
    r = para.add_run("")._r
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), "begin")
    if dirty is not None:
        fld.set(qn('w:dirty'), "true" if dirty else "false")
    r.append(fld)

    r = para.add_run("")._r
    cmd = OxmlElement('w:instrText')
    cmd.text = instr
    r.append(cmd)

    r = para.add_run("")._r
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), "separate")
    r.append(fld)

    if cache is not None:
        para = cache(para)

    r = para.add_run("")._r
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), "end")
    r.append(fld)


def docx_add_caption(para: Paragraph, title: Text, caption: Text  # {{{1
                     ) -> None:
    """
        <w:p><w:pPr>
            <w:pStyle w:val="Normal"/><w:bidi w:val="0"/><w:jc w:val="left"/>
            <w:rPr></w:rPr></w:pPr>
          <w:r><mc:AlternateContent><mc:Choice Requires="wps">
            <w:drawing>
            <wp:anchor behindDoc="0" distT="0" distB="0" distL="0"
                       distR="0" simplePos="0" locked="0" layoutInCell="1"
                       allowOverlap="1" relativeHeight="2">
              <wp:simplePos x="0" y="0"/>
              <wp:positionH relativeFrom="column">
                <wp:align>center</wp:align></wp:positionH>
              <wp:positionV relativeFrom="paragraph">
                <wp:posOffset>635</wp:posOffset></wp:positionV>
              <wp:extent cx="5409565" cy="1374140"/>
              <wp:effectExtent l="0" t="0" r="0" b="0"/>
              <wp:wrapSquare wrapText="largest"/>
              <wp:docPr id="1" name="Frame1"/>
              <a:graphic
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData
       uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                  <wps:wsp><wps:cNvSpPr txBox="1"/>
                    <wps:spPr>
                      <a:xfrm><a:off x="0" y="0"/>
                        <a:ext cx="5409565" cy="1374140"/>
                      </a:xfrm>
                      <a:prstGeom prst="rect"/>
                    </wps:spPr>
                    <wps:txbx><w:txbxContent>
                      <w:p>
                        <w:pPr><w:pStyle w:val="Figure"/>
                               <w:bidi w:val="0"/>
                               <w:spacing w:before="120" w:after="120"/>
                               <w:jc w:val="left"/>
                               <w:rPr></w:rPr>
                        </w:pPr>
                        <w:r><w:drawing>
                          <wp:inline distT="0" distB="0" distL="0" distR="0">
                          <wp:extent cx="5409565" cy="1089660"/>
                          <wp:effectExtent l="0" t="0" r="0" b="0"/>
                          <wp:docPr id="2" name="Image1" descr=""></wp:docPr>
                          <wp:cNvGraphicFramePr>
                            <a:graphicFrameLocks
                xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                                                           noChangeAspect="1"/>
                        </wp:cNvGraphicFramePr>
                        <a:graphic
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                                <a:graphicData
                uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                                    <pic:pic
          xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                          <pic:nvPicPr>
                            <pic:cNvPr id="2" name="Image1" descr="">
                            </pic:cNvPr>
                            <pic:cNvPicPr>
                              <a:picLocks noChangeAspect="1"
                                          noChangeArrowheads="1"/>
                            </pic:cNvPicPr>
                          </pic:nvPicPr>
                          <pic:blipFill>
                            <a:blip r:embed="rId2"></a:blip>
                            <a:stretch><a:fillRect/></a:stretch>
                          </pic:blipFill>
                          <pic:spPr bwMode="auto">
                            <a:xfrm><a:off x="0" y="0"/>
                                    <a:ext cx="5409565" cy="1089660"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                          </pic:spPr>
                        </pic:pic></a:graphicData> </a:graphic>
                        </wp:inline></w:drawing>
                        </w:r>
            <w:r><w:rPr><w:vanish/></w:rPr><w:br/></w:r>

            <!-- core part -->
            <w:r><w:t xml:space="preserve">Figure </w:t></w:r>
            <w:r><w:fldChar w:fldCharType="begin"></w:fldChar></w:r>
            <w:r><w:instrText> SEQ Figure \* ARABIC</w:instrText></w:r>
            <w:r><w:fldChar w:fldCharType="separate"/></w:r>
            <w:r><w:t>1</w:t></w:r>
            <w:r><w:fldChar w:fldCharType="end"/></w:r>
            <w:r><w:t>.  aaa</w:t></w:r>

            </w:p>
          </w:txbxContent></wps:txbx>
          <wps:bodyPr anchor="t" lIns="0" tIns="0" rIns="0" bIns="0">
            <a:noAutofit/></wps:bodyPr>
        </wps:wsp>
        </a:graphicData></a:graphic></wp:anchor>
        </w:drawing></mc:Choice><mc:Fallback>
        <w:pict><v:rect style="position:absolute;rotation:0;width:425.95pt;
                               height:108.2pt;mso-wrap-distance-left:0pt;
                               mso-wrap-distance-right:0pt;
                               mso-wrap-distance-top:0pt;
                               mso-wrap-distance-bottom:0pt;
                               margin-top:0pt;
                               mso-position-vertical:top;
                               mso-position-vertical-relative:text;
                               margin-left:28pt;
                               mso-position-horizontal:center;
                               mso-position-horizontal-relative:text">
          <v:textbox...
    """
    if caption not in glb.numbers_of_captions:
        glb.numbers_of_captions[caption] = 1
    else:
        glb.numbers_of_captions[caption] += 1
    n = glb.numbers_of_captions[caption]

    def cb(p: Paragraph) -> Paragraph:
        p.add_run("%d" % n)
        return p

    # TODO(shimoda): surround figure by frame
    para.add_run(caption + " ")  # r:vanish??
    docx_add_field(para, r"SEQ %s \* ARABIC" % caption, cb)
    para.add_run(". " + title)


def docx_bookmark_normalize(name: Text) -> Text:  # {{{1
    if len(name) > 32:
        # [@P8-1-1]
        pfx, sfx = name[:32], name[32:]
        hash = zlib.adler32(name.encode("utf-8")) & 0xFFFFFFFF
        sfx = ("0" * 8 + hex(hash)[2:])[-8:]
        name = pfx + sfx
    return name


def docx_add_bookmark(para: Paragraph, name: Text, instr: Text):  # {{{1
    id_ = glb.bookmark_id

    def mark(tag: Text, name: Text):
        mk = OxmlElement("w:bookmark" + tag)
        mk.set(qn("w:id"), "%d" % id_)
        if len(name) > 0:
            name = docx_bookmark_normalize(name)
            mk.set(qn("w:name"), name)
            debg("book: " + name)
        para._p.append(mk)

    if len(name) > 0:
        mark("Start", name)

    para.add_run(instr)

    if len(name) > 0:
        mark("End", "")
        glb.bookmark_id += 1


def docx_add_hlink(para: Paragraph, instr: Text,  # {{{1
                   name: Text) -> None:
    link = OxmlElement("w:hyperlink")
    para._p.append(link)
    run = OxmlElement("w:r")
    link.append(run)
    text = OxmlElement("w:t")
    run.append(text)

    name = docx_bookmark_normalize(name)
    link.set(qn("w:anchor"), name)
    debg("link: " + name)
    text.text = instr


def style(name: Text  # {{{1
          ) -> Callable[[Callable[['Styles', Document, Text], Text]],
                        Callable[['Styles', Document], Text]]:

    def ret(fn: Callable[['Styles', Document, Text], Text]
            ) -> Callable[['Styles', Document], Text]:

        def new_fn(self: 'Styles', doc: Document) -> Text:
            return fn(self, doc, name)

        glb.seq_styles[name] = new_fn
        return new_fn

    return ret


class Styles(object):  # {{{1
    @classmethod
    def get(cls, doc: Document, name: Text) -> Text:
        if name not in glb.seq_styles:
            raise ParseError("")
        fn = glb.seq_styles[name]
        self = Styles()
        return fn(self, doc)

    def init_heading(self, doc: Document, name: Text) -> Text:  # {{{1
        # failure code.
        # st = doc.styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
        # st.base_style = doc.styles["Heading 1"]
        # st.font.color.rgb = RGBColor(0, 0, 0)
        doc.styles[name].font.color.rgb = RGBColor(0, 0, 0)
        return name

    for i in range(1, 10):
        (style("Heading %d" % i))(init_heading)

    @style("Quote")  # {{{1
    def init3(self, doc: Document, name: Text) -> Text:
        style = doc.styles['Quote']
        style.font.name = "SourceCodePro"
        style.font.size = Pt(8)
        style.font.italic = False
        fmt = doc.styles['Quote'].paragraph_format
        fmt.line_spacing = Pt(9)
        fmt.left_indent = fmt.right_indent = Mm(10)
        return name

    @classmethod  # quote {{{1
    def quote(self, para: Paragraph) -> None:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        for val in ["left", "right", "top", "bottom"]:
            b = OxmlElement('w:' + val)
            b.set(qn('w:val'), 'thinThickLargeGap')
            b.set(qn('w:sz'), '2')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), '000000')
            pBdr.append(b)

    @style("CodeChars")  # {{{1
    def init_codechars(self, doc: Document, name: Text) -> Text:
        fmt = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        fmt.font.name = "SourceCodePro"
        fmt.font.size = Pt(8)
        fmt.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return name

    @style("katex")  # {{{1
    def init_katex(self, doc: Document, name: Text) -> Text:
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        return name

    @style("Caption")  # {{{1
    def init_caption(self, doc: Document, name: Text) -> Text:
        fmt = doc.styles['Caption'].paragraph_format
        fmt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return name

    @style("Image")  # {{{1
    def init_image(self, doc: Document, name: Text) -> Text:
        "[@P14-1-11] style for single images"
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return name

    def init_list(self, doc: Document, name: Text) -> Text:  # {{{1
        # list items indent
        if name[-1] in "0123456789":
            n = int(name.split(" ")[-1])
        else:
            n = 1
        if name in doc.styles:
            style = doc.styles[name]
        else:
            org = " ".join(name.split(" ")[:-1])
            style_org = doc.styles[org]
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_org
        fmt = style.paragraph_format
        fmt.left_indent = Mm(15 + 15 * (n - 1))
        fmt.first_line_indent = Mm(-5)
        # [@P11-1-1] reduce extra spaces after paragraph
        fmt.space_after = Mm(1)
        fmt.line_spacing = 1.0
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

        """ failed to set
        ppr= i.element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ppr.append(ind)
        ind.set(qn("w:left"), "%d" % Mm(25 + 20 * n))      # - <---hanging---|
        ind.set(qn("w:hanging"), "%d" % Mm(10 + 20 * n))   # ------left----->|
        """
        return name

    (style("List Number"))(init_list)
    (style("List Bullet"))(init_list)
    for i in range(2, 10):
        (style("List Number %d" % i))(init_list)
        (style("List Bullet %d" % i))(init_list)

    def init_toc(self, doc: Document, name: Text) -> Text:  # {{{1
        # [@P5-1-11] line spacing of TOC
        level = name.split(" ")[2]
        sname = "List" + ((" " + level) if level != "1" else "")
        style = doc.styles.add_style(
                name, WD_STYLE_TYPE.PARAGRAPH
                )
        style.base_style = doc.styles[sname]
        fmt = style.paragraph_format
        fmt.space_after = Mm(1)
        fmt.line_spacing = 1.0
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        return name

    for i in range(1, 10):
        (style("TOC Contents %d" % i))(init_toc)

    @style("Stamps")  # {{{1
    def init_stamps(self, doc: Document, name: Text) -> Text:
        # styles for stamps
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(8)
        return name

    @style("CellHeader")  # {{{1
    def init_cell_header(self, doc: Document, name: Text) -> Text:
        "[@P13-2-11] style for header cells"
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return name

    @style("CellNormal")  # {{{1
    def init_cell_normal(self, doc: Document, name: Text) -> Text:
        "[@P13-2-12] style for cells"
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        return name

    @style("Subtitle")  # {{{1
    def init_subtitle(self, doc: Document, name: Text) -> Text:
        style = doc.styles[name]
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)
        return name

    @style("Title")  # {{{1
    def init_title(self, doc: Document, name: Text) -> Text:
        style = doc.styles[name]
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style.font.size = Pt(20)
        style.font.color.rgb = RGBColor(0, 0, 0)
        ppr = style.element.get_or_add_pPr()   # remove border
        seq = ppr.xpath("w:pBdr")
        if len(seq) > 0:
            ppr.remove(seq[0])
        return name


def docx_style(doc: Document, name: Text) -> Text:  # {{{1
    if name in glb.styles:
        return name

    try:
        name = Styles.get(doc, name)
        glb.styles.append(name)
        return name
    except ParseError:
        pass

    for style in doc.styles:
        if name != style.name:
            continue
        glb.styles.append(name)
        return name
    assert False, "can not found style: %s" % name


def main(args: List[Text]) -> int:  # {{{1
    fname = args[1]
    size = image_width_and_height(fname)
    print("%s->%s" % (fname, size))
    size2 = tuple(dot_to_mm(i) for i in size)
    print("%s->%s" % (fname, size2))
    argx = dot_to_page(size2[0], size2[1])
    print("%s->%s" % (fname, argx))
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))

# {{{1 end of file
# vi: ft=python:fdm=marker
