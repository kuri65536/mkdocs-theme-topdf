# Copyright (c) 2020, Shimoda <kuri65536 at hotmail dot com>
#
# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at http://mozilla.org/MPL/2.0/.
#
# include {{{1
import logging
from logging import (debug as debg, info, warning as warn, )
from lxml import etree  # type: ignore
import os
import re
from typing import (Dict, Iterable, List, Optional, Text, Tuple, Union, )

from bs4 import BeautifulSoup  # type: ignore
from bs4.element import Tag  # type: ignore

from docx import Document  # type: ignore
from docx.blkcntnr import BlockItemContainer  # type: ignore
from docx.enum.text import (                  # type: ignore
        WD_ALIGN_PARAGRAPH, WD_BREAK,  # type: ignore
        )             # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.shared import Mm  # type: ignore
from docx.table import _Cell, Table  # type: ignore

try:
    from . import common
    from . import options
    from . import docx_toc
    from . import docx_svg_hack
except ImportError:
    import common  # type: ignore
    import options  # type: ignore
    import docx_toc  # type: ignore
    import docx_svg_hack  # type: ignore


if False:
    List


class _info_list:
    def __init__(self, f: bool, s: Text, l: int) -> None:
        self.f_number = f
        self.style = s
        self.level = l


class HtmlConvertDocx(object):  # {{{1
    def __init__(self, src: Text) -> None:  # {{{1
        self.bookmarks_anchored: Dict[Text, Text] = {}
        if common.is_target_in_http(src):
            self.url_target = src
        else:
            src = os.path.realpath(src)
            self.url_target = src

        doc = Document()
        self.output = doc
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        doc.settings.element.append(update_fields)
        info("structure root")
        self.para: Optional[Paragraph] = None  # doc.add_paragraph('')
        self.header_init()

    def header_init(self) -> None:  # {{{1
        sec = self.output.sections[0]

        # page setting
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = sec.right_margin = sec.bottom_margin = Mm(20)
        sec.top_margin = Mm(20)  # 20 - 1

        # make header to right
        sec.header_distance = Mm(13)  # 20 - 1
        para = self.output.sections[0].header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # page border as rectangle (old word style)
        """<w:pict><v:rect id="shape_0" ID="Shape1" stroked="t"
                style="position:absolute;margin-left:-5.8pt;margin-top:13.05pt;width:493.2pt;height:751.15pt">
                <w10:wrap type="none"/>
                <v:fill o:detectmouseclick="t" on="false"/>
                <v:stroke color="black" weight="12600" joinstyle="round"
                          endcap="flat"/>
           </v:rect></w:pict>
        """
        v = "urn:schemas-microsoft-com:vml"
        w10 = "urn:schemas-microsoft-com:office:word"
        r = para._p.add_r()
        pict = OxmlElement('w:pict')
        r.append(pict)
        rect = etree.Element("{%s}rect" % v)
        pict.append(rect)
        rect.set('id', 'shape_0')  # page_border')
        rect.set('ID', 'Shape1')   # _page_border')
        rect.set('stroked', 't')
        rect.set('style', 'position:absolute;'
                 'margin-left:-5.8pt;margin-top:20.0pt;'
                 'width:493.2pt;height:751.15pt')
        wrap = etree.Element('{%s}wrap' % w10)
        wrap.set('type', "none")
        rect.append(wrap)
        fill = etree.Element('{%s}fill' % v)
        fill.set('on', "false")
        rect.append(fill)
        stroke = etree.Element('{%s}stroke' % v)
        stroke.set('color', "black")
        stroke.set('weight', "12600")
        stroke.set('joinstyle', "round")
        stroke.set('endcap', "flat")
        rect.append(stroke)

    def header_set(self, src: Text) -> None:  # {{{1
        para = self.output.sections[0].header.paragraphs[0]
        para.add_run(src + "( ")
        common.docx_add_field(para, "PAGE", None)
        para.add_run(" / ")
        common.docx_add_field(para, "NUMPAGES", None)
        para.add_run(" )")

    def write_out(self, fname: Text) -> None:
        info("structure save")
        self.output.save(fname)

        doc = Document(fname)
        for para in doc.paragraphs:
            ret = Text(para.text)
            info("after-para: " + (ret[0:10] if len(ret) > 9 else ret))
        for tbl in doc.tables:
            ret = "%d,%d" % (len(tbl.rows), len(tbl.rows[0].cells))
            info("after-tabl: " + (ret))

    def on_post_page(self, output_content: Text,  # {{{1
                     backend_bs4: Text) -> Text:
        soup = BeautifulSoup(output_content, backend_bs4)
        dom = soup.find("body")
        self.extract_para(dom, 0)
        return output_content

    def extract_is_text(self, elem: Tag) -> Union[None, Text, Tag]:  # {{{1
        if elem.name is not None:
            return elem
        elif "element.Comment" in Text(type(elem)):
            return None  # ignore html comments
        return self.extract_text(elem)

    def extract_inlines(self, elem: Tag, para: Paragraph  # {{{1
                        ) -> Optional[Text]:
        # inline elements
        if elem.name == "em":
            return self.extract_em(elem, para)
        elif elem.name == "strong":
            return self.extract_strong(elem, para)
        elif elem.name == "sup":
            return self.extract_sup(elem, para)
        elif elem.name == "code":
            return self.extract_code(elem, para, pre=False)
        elif elem.name == "br":
            return self.extract_br(elem, para)
        elif elem.name == "a":
            return self.extract_anchor(elem, para)
        elif elem.name == "span" and common.has_class(elem, "katex-display"):
            return self.extract_katex(elem, para)
        elif elem.name == "span":
            text = elem.text
            para = self.current_para_or_create(para)
            para.add_run(text)
            return text
        raise common.ParseError("%s is not implemented, yet" % elem.name)

    def extract_element(self, elem: Tag, para: Paragraph  # {{{1
                        ) -> Union[None, Text, Tag]:
        is_text = self.extract_is_text(elem)
        if (is_text is None) or isinstance(is_text, Text):
            return is_text

        classes = elem.attrs.get("class", [])
        if "doc-num" in classes:
            self.header_set(elem.string)
            return None
        if "toc" in classes:
            self.para = docx_toc.generate_toc(self.output, elem)
            return None

        try:
            return self.extract_inlines(elem, para)
        except common.ParseError:
            pass

        # block elements
        if elem.name in ("script", "style"):
            return None  # just ignore these elements.
        elif elem.name == "hr":
            return self.extract_pagebreak(elem)
        elif elem.name == "dl":
            return self.extract_dldtdd(elem)
        elif elem.name == "ul":
            return self.extract_list(elem, False, 1, self.output)
        elif elem.name == "ol":
            return self.extract_list(elem, True, 1, self.output)
        elif elem.name == "table":
            return self.extract_table(elem)
        elif elem.name == "details":
            return self.extract_details(elem)
        elif elem.name == "img":
            return self.extract_img(elem, para)
        elif elem.name == "svg":
            return self.extract_svg(elem, para)
        elif elem.name in ("h1", "h2", "h3", "h4", "h5", "h6", ):
            return self.extract_title(elem)
        elif elem.name in ("pre", "code"):
            return self.extract_codeblock(elem)
        elif elem.name in ("p", "div"):
            pass
        else:
            pass  # for debug: import pdb; pdb.set_trace()
        # p, article or ...
        return elem

    def extract_text(self, elem: Tag) -> Optional[Text]:  # {{{1
        if elem.string == "\n":
            if elem.parent.name in ("body", "div", ):
                return None
        ret = Text(elem.string)
        # [@D4-19-1] treat \n as empty characters.
        ret = re.sub(r" *\n *", " ", ret)
        debg(ret.strip())
        return ret

    def extract_br(self, elem: Tag, para: Paragraph) -> Text:  # {{{1
        # sometime bs4 fails to parse `br` tag and surround text.
        n, ret = 0, ""
        for tag in elem.children:
            n += 1
            content = self.extract_is_text(tag)
            if content is None:
                continue
            if isinstance(content, Text):
                para.add_run(content)
            else:
                content = self.extract_inlines(tag, para)
                if not isinstance(content, Text):
                    continue
            ret += content
        if n < 1:
            if para is None:
                para = self.output.add_paragraph()
            para.add_run("\n")
        return ret

    def extract_em(self, elem: Tag, para: Paragraph) -> Optional[Text]:  # {{{1
        def add_para() -> Paragraph:
            style = common.docx_style(self.output, "Caption")
            self.para = self.output.add_paragraph("", style)
            return self.para
        classes = elem.attrs.get("class", [])
        if "table-tag" in classes:
            common.docx_add_caption(add_para(), elem.text, "Table")
            return None
        if common.has_prev_sibling(elem, "img", "svg"):
            common.docx_add_caption(add_para(), elem.text, "Figure")
            return None
        if common.has_next_table(elem):
            common.docx_add_caption(add_para(), elem.text, "Table")
            return None

        para = self.current_para_or_create(para)
        para.add_run(elem.text, style="Emphasis")
        return None

    def extract_code(self, elem: Tag, para: Paragraph, pre: bool  # {{{1
                     ) -> Optional[Text]:
        s = Text(elem.text)
        if para is None:  # top level
            style = common.docx_style(self.output, "Quote")
            self.output.add_paragraph(s, style)
        elif pre:
            para.add_run(s)
        else:
            style = common.docx_style(self.output, "CodeChars")
            para.add_run(s, style=style)
        return None

    def extract_strong(self, elem: Tag, para: Paragraph  # {{{1
                       ) -> Optional[Text]:
        s = Text(elem.text)
        style = common.docx_style(self.output, "Strong")
        para = self.current_para_or_create(para)
        para.add_run(s, style=style)
        return None

    def extract_sup(self, elem: Tag, para: Paragraph  # {{{1
                    ) -> Optional[Text]:
        # [@D4-99-001-1] treat sup elements
        ret = ""
        for tag in elem.children:
            if tag.name == "a":
                para = self.current_para_or_create(para)
                para.add_run(ret)
                ret = ""
                # TODO(shimoda): append extra 'sup' styles for anchor.
                self.extract_anchor(tag, para)
                continue

            content = self.extract_is_text(tag)
            if isinstance(content, Text):
                ret += content
            elif content is not None:
                breakpoint()
        if len(ret) > 0:
            para = self.current_para_or_create(para)
            para.add_run(ret)
        return None

    def extract_anchor(self, elem: Tag, para: Paragraph  # {{{1
                       ) -> Optional[Text]:
        instr = Text(elem.text)
        url = elem.attrs.get("href", "")
        if len(url) < 1:
            name = elem.attrs.get("name", "")
            if len(name) < 1:
                return instr
            if self.bookmark_from_db(name, elem):
                return instr
            common.docx_add_bookmark(para, "ahref_" + name, instr)
            return None
        if not url.startswith("#"):
            # TODO(shimoda): enable external link
            return instr

        para = self.current_para_or_create(para)
        # remove "#"
        common.docx_add_hlink(para, instr, "ahref_" + url[1:])
        return None

    def extract_katex(self, elem: Tag, para: Paragraph  # {{{1
                      ) -> Optional[Text]:
        """html sample: moved to test/docs/report-docx.md
        """
        # TODO(shimoda): convert to image?
        anno = elem.find("annotation")
        if anno is not None:
            style = common.docx_style(self.output, "katex")
            self.para = self.output.add_paragraph(anno.text, style)
        else:
            breakpoint()
        return None

    def extract_table_tree(self, elem: Tag, row: int  # {{{1
                           ) -> Dict[Tuple[int, int], Tag]:
        def num_row(dct: Dict[Tuple[int, int], Tag]) -> int:
            ret = -1
            for (row, col) in dct.keys():
                ret = max(ret, row)
            return ret + 1

        ret: Dict[Tuple[int, int], Tag] = {}
        f_row = False
        for tag in elem.children:
            if tag.name is None:
                continue
            if tag.name == "thead":
                info("structure: tbl: enter thead")
                ret = self.extract_table_tree(tag, 0)
                continue
            if tag.name == "tbody":
                info("structure: tbl: enter tbody")
                ret2 = self.extract_table_tree(tag, num_row(ret))
                ret.update(ret2)
                continue
            if tag.name == "tfoot":
                warn("structure: tbl: enter tfoot")
                ret2 = self.extract_table_tree(tag, num_row(ret))
                ret.update(ret2)
                continue
            if tag.name != "tr":
                warn("table tree: %s in table" % tag.name)
                continue
            f_row, col, row = True, 0, (row + 1 if f_row else row)
            for cell in tag.children:
                if cell.name is None:
                    continue
                if cell.name not in ("th", "td", ):
                    warn("table tree: %s in tr" % cell.name)
                    continue
                ret[row, col] = cell
                col += 1
        return ret

    def extract_table(self, elem: Tag) -> Optional[Text]:  # {{{1
        dct = self.extract_table_tree(elem, 0)
        dct = common.table_update_rowcolspan(dct)  # [@P13-1-11] cell-span
        if len(dct) < 1:
            warn("table: did not have any data" + Text(elem.string))
            return None
        n_row = max([tup[0] for tup in dct.keys()]) + 1
        n_col = max([tup[1] for tup in dct.keys()]) + 1

        info("structure: tbl: %s (%d,%d)" % (elem.name, n_row, n_col))
        tbl = self.output.add_table(rows=n_row, cols=n_col)
        tbl.autofit = False
        tbl.style = common.docx_style(self.output, "Table Grid")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for (row, col), subelem in sorted(dct.items(), key=lambda x: x[0]):
            cell = tbl.rows[row].cells[col]
            # [@P13-1-12] merging spaned cells
            x, y = common.table_cellspan(subelem, "colspan", "rowspan")
            if x != 0 or y != 0:
                if col + x >= n_col or row + y >= n_row:
                    raise IndexError("invalid colspan or rowspan in '%s'" %
                                     Text(elem))
                cel2 = tbl.rows[row + y].cells[col + x]
                cell.merge(cel2)
            self.extract_table_cell(subelem, cell)

        classes = common.classes_from_prev_sibling(elem)
        self.style_table_width_from(tbl, classes)
        return None

    def extract_pagebreak(self, elem: Tag) -> Optional[Text]:  # {{{1
        if self.para is None:
            self.para = self.output.add_paragraph()
        self.para.add_run().add_break(WD_BREAK.PAGE)
        return None

    def extract_dldtdd(self, elem: Tag) -> Optional[Text]:  # {{{1
        classes = common.classes_from_prev_sibling(elem)
        if "before-dl-table" not in classes:
            warn("can not convert dl tags, "
                 "docx does not have difinition lists")
            return None

        n_row = n_col = i = 0
        for tag in elem.children:
            if tag.name not in ("dt", "dd"):
                continue
            if tag.name == "dt":
                n_row, i = n_row + 1, 1
                continue
            i += 1
            n_col = max(i, n_col)

        info("structure: tbl: %s (%d,%d)" % (elem.name, n_row, n_col))
        tbl = self.output.add_table(rows=n_row, cols=n_col)
        tbl.autofit = False
        tbl.style = common.docx_style(self.output, "Table Grid")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        n_row, i = -1, 0
        for tag in elem.children:
            if tag.name not in ("dt", "dd"):
                continue
            elif tag.name == "dt":
                n_row, i = n_row + 1, 0
            else:
                i += 1
            self.extract_table_cell(tag, tbl.rows[n_row].cells[i])
        if self.style_table_stamps(tbl, classes):
            return None
        self.style_table_width_from(tbl, classes)
        return None

    def extract_list(self, elem: Tag, f_number: bool,  # {{{1
                     level: int, blk: BlockItemContainer) -> Optional[Text]:
        list_info = self.style_list(f_number, level)
        for tag in elem.children:
            if tag.name != "li":
                continue
            ret = self.extract_list_subs(None, tag, list_info, blk)
            info("structure: li : " + ret if len(ret) < 50 else ret[:50])
        self.para = None
        return None

    def extract_details(self, elem: Tag) -> Optional[Text]:
        debg("structure: details: not implemented, skipped...")
        return None

    def extract_img(self, elem: Tag, para: Paragraph  # {{{1
                    ) -> Optional[Text]:
        para = self.current_para_or_create(para)
        # [@P14-1-11] center images at some rules.
        if common.count_tags_around_image(elem.parent) <= 2:
            para.style = common.docx_style(self.output, "Image")

        src = elem.attrs.get("src", "")
        if len(src) < 1:
            warn("img-tag: was not specified 'src' attribute, ignored...")
            return None
        fname = common.download_image(self.url_target, src)
        if len(fname) < 1:
            warn("img-tag: can not download, ignored...: " + src)
            return None

        w, h = common.image_width_and_height(fname)
        if (w, h) == (0, 0):
            # TODO(shimoda): implement for svg
            return None
        args = common.dot_to_page(w, h)
        para.add_run().add_picture(fname, **args)
        return None

    def extract_svg(self, elem: Tag, para: Paragraph  # {{{1
                    ) -> Optional[Text]:
        # [P10-1-11] convert to an imported svg.
        fname = docx_svg_hack.dump_file(elem, "tmp")
        docx_svg_hack.monkey()

        para = self.current_para_or_create(para)
        run = para.add_run()
        pic = run.add_picture(fname)

        docx_svg_hack.compose_asvg(run, pic)
        return None

    def extract_title(self, elem: Tag) -> Optional[Text]:  # {{{1
        level = int(elem.name.lstrip("h"))
        ret = Text(elem.text)
        info("structure: hed: " + ret)
        style = common.docx_style(self.output, "Heading " + Text(level))
        para = self.output.add_paragraph("", style=style)

        # [@P8-2-13] add id for titles
        html_id = elem.attrs.get("id", "")
        common.docx_add_bookmark(para, "ahref_" + html_id, ret)
        self.para = None
        return None

    def extract_codeblock(self, elem: Tag) -> Optional[Text]:  # {{{1
        ret = Text(elem.string)
        info("structure: pre: " + ret.splitlines()[0])
        style = common.docx_style(self.output, "Quote")
        para = self.output.add_paragraph(ret, style=style)
        self.para = None
        common.Styles.quote(para)
        return None

    def extract_para(self, node: Tag, level: int) -> Optional[Text]:  # {{{1
        info("enter para...: %d-%s" % (level, node.name))
        if (node.name == "p" and
                common.has_class(node, *options.current.classes_ignore_p)):
            return None
        bkname = self.bookmark_from_elem(node)  # [@P8-2-14] mark for <p>

        para = None
        for elem in node.children:
            ret = self.extract_element(elem, para)
            if isinstance(ret, Text):
                empty = ret.strip("\n")
                if len(empty) < 1:
                    pass
                elif para is None and len(bkname) > 0:  # [@P8-2-15] mark for p
                    para = self.current_para_or_create(para)
                    common.docx_add_bookmark(para, bkname, ret)
                elif para is None:
                    self.para = para = self.output.add_paragraph(ret)
                else:
                    para.add_run(ret)
            elif ret is not None:
                self.extract_para(elem, level + 1)
            else:
                if para != self.para:
                    para = self.para
        return None

    def extract_table_cell(self, elem: Tag, cell: _Cell) -> None:  # {{{1
        # FIXME(shimoda): simplify complex code...
        para: Optional[Paragraph] = cell.paragraphs[-1]
        for tag in elem.children:
            src = ""
            try:
                self.extract_inlines(tag, para)
                continue
            except common.ParseError:
                pass
            if tag.name is None:
                src = tag.string
                src = src.replace("\n", " ")
            elif tag.name in ("p", "div"):
                if para is None:
                    para = cell.add_paragraph()
                self.extract_as_run(para, tag, "")
            elif tag.name in ("ul", "ol"):
                self.extract_list(tag, tag.name == "ol", 1, cell)
                para = None
                continue
            else:
                debg("can't extract %s in a table-cell" % tag.name)
                continue
            if para is None:
                para = cell.add_paragraph()
            info("table: %s" % src)  # (row, col, src))
            para.add_run(src)

        # [@P13-2-13] style for cells
        style_name = "CellHeader" if elem.name == "th" else "CellNormal"
        style = common.docx_style(self.output, style_name)
        for para in cell.paragraphs:
            para.style = style

    def extract_list_subs(self, para: Optional[Paragraph], elem: Tag,  # {{{1
                          info: _info_list, blk: BlockItemContainer) -> Text:
        ret = ""
        for tag in elem.children:
            if tag.name in ("p", "div", "ul", "ol"):
                break
        else:
            if para is None:
                para = blk.add_paragraph("", info.style)
            bkname = self.bookmark_from_elem(elem)  # [@P8-2-11] mark for li-1
            return self.extract_as_run(para, elem, bkname)
        for tag in elem.children:
            if tag.name in ("ul", "ol"):
                self.extract_list(tag, tag.name == "ol", info.level + 1, blk)
                para = None
                continue
            elif tag.name in ("p", "div"):
                ret += self.extract_list_subs(para, tag, info, blk)
                continue

            try:
                self.extract_inlines(tag, para)
                continue
            except common.ParseError:
                pass
            if tag.name is None:
                src = tag.string
                src = src.replace("\n", " ")
                if len(src.strip()) > 0:
                    if para is None:
                        para = blk.add_paragraph("", info.style)
                    para.add_run(src)
                    ret += src
            else:
                warn("can't parse complex html...%s" % tag.name)
        return ret

    def extract_as_run(self, para: Paragraph, elem: Tag,  # {{{1
                       bkname: Text) -> Text:
        ret = ""
        if elem.name is None:  # NavigableString
            ret = Text(elem.string)
            common.docx_add_bookmark(para, bkname, ret)  # [@P8-2-11] for li
            return ret
        for tag in elem.children:
            content = self.extract_is_text(tag)
            if isinstance(content, Text):
                # [@P8-2-12] for li, head of paragraph
                ret += content
                common.docx_add_bookmark(para, bkname, content)
                bkname = ""
                continue
            if content is None:
                continue
            try:
                s = self.extract_inlines(
                        tag, para)
                if isinstance(s, Text):
                    ret += s
                if len(bkname) > 0:
                    common.docx_add_bookmark(para, bkname, " ")  # [@P8-2-13]
            except common.ParseError:
                pass
        return ret

    def style_table_stamps(self, tbl: Table,  # {{{1
                           classes: Iterable[Text]
                           ) -> bool:
        if "table-3stamps" not in classes:
            return False
        for col, wid in zip(tbl.columns,
                            [Mm(114), Mm(20), Mm(20), Mm(20)]):
            col.cells[0].width = wid
        row = tbl.rows[0]
        row.height = Mm(20)
        mar = OxmlElement('w:tblCellMar')
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.allow_autofit = True
        tbl._tblPr.append(mar)
        for ename, val in (('top', '0'), ('bottom', '0'),
                           ('left', '108'), ('right', '108'), ):
            i = OxmlElement('w:' + ename)
            i.set(qn('w:w'), val)
            i.set(qn('w:type'), 'dxa')
            mar.append(i)
        i = OxmlElement('w:tblInd')
        i.set(qn('w:w'), "-6")
        i.set(qn('w:type'), 'dxa')
        mar.append(i)
        """
            <w:tbl><w:tblPr>
              <w:tblStyle w:val="TableGrid"/>
              <w:tblW w:w="9870" w:type="dxa"/>
              <w:jc w:val="left"/>
              <w:tblInd w:w="-6" w:type="dxa"/>
              <w:tblCellMar>
                <w:top w:w="0" w:type="dxa"/>
                <w:left w:w="108" w:type="dxa"/>
                <w:bottom w:w="0" w:type="dxa"/>
                <w:right w:w="108" w:type="dxa"/>
              </w:tblCellMar>
              <w:tblLook w:lastRow="0"
                  w:firstRow="1" w:lastColumn="0" w:firstColumn="1"
                  ...
        """

        n = 0
        for para in row.cells[0].paragraphs:
            n += 1
        if n == 1:
            # style up <dt>subtitle<br>title</dt>
            para = row.cells[0].paragraphs[-1]
            lines = para.text.split("\n")
            if len(lines) < 2:
                para.style = common.docx_style(self.output, "Title")
                return True
            para.text = lines[0]
            para = row.cells[0].add_paragraph("\n".join(lines[1:]))

        for n, para in enumerate(row.cells[0].paragraphs):
            if n == 0:
                para.style = common.docx_style(self.output, "Subtitle")
            else:
                para.style = common.docx_style(self.output, "Title")
        for i in range(1, 4):
            row.cells[i].paragraphs[0].style = common.docx_style(
                    self.output, "Stamps")
            if "\n" not in row.cells[i].text:
                tcpr = row.cells[i]._element.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                tcpr.append(borders)
                bs = OxmlElement("w:tl2br")
                borders.append(bs)
                bs.set(qn("w:color"), "000000")
                bs.set(qn("w:space"), "0")
                bs.set(qn("w:val"), "single")
                bs.set(qn("w:sz"), "4")  # 1pt
                # can not set tl2br in LibreOffice.
                # this code was confirmed by w:bottom and color 00FF00
        return True

    def style_table_width_from(self, tbl: Table,  # {{{1
                               classes: Iterable[Text]
                               ) -> bool:
        # total width: 160mm
        cls, widths = common.has_width_and_parse(classes)
        if len(widths) < 1:
            warn("width did not specified by class")
            return False
        warn("cell width set by %s" % cls)
        if Mm(0) in widths:
            tbl.autofit = True
            tbl.allow_autofit = True
        for j, row in enumerate(tbl.rows):
            for i, cell in enumerate(row.cells):
                wid = widths[i] if i < len(widths) else 0
                if wid < 1:
                    info("cell(%d,%d): width set to auto" % (j, i))
                    cell._tc.tcPr.tcW.type = 'auto'
                    cell._tc.tcPr.tcW.w = 0
                    continue
                info("cell(%d,%d): width set to %d" % (j, i, wid))
                cell.width = wid
        return True

    def style_list(self, f_number: bool, level: int) -> _info_list:  # {{{1
        style_base = "List Number" if f_number else "List Bullet"
        if level > 1:
            style = style_base + " %d" % level
        else:
            style = style_base
        style = common.docx_style(self.output, style)
        return _info_list(f_number, style, level)

    def current_para_or_create(self, para: Paragraph, style: Text = ""  # {{{1
                               ) -> Paragraph:
        if para is not None:
            return para
        if len(style) < 1:
            para = self.para = self.output.add_paragraph()
            return para
        para = self.para = self.output.add_paragraph("", style)
        return para

    def bookmark_from_elem(self, elem: Tag) -> Text:  # {{{1
        # [@P8-2-14]
        id_ = elem.attrs.get("id", "")
        if len(id_) < 1:
            return ""
        if self.bookmark_from_db(id_, elem):
            return ""  # this id is not anchored -> do not convert.
        return "ahref_" + id_

    def bookmark_from_db(self, s_href: Text, elem: Tag) -> bool:  # {{{1
        db = self.bookmarks_anchored
        if len(db) < 1:
            seq = [i for i in elem.parents if i.name == "body"]
            if len(seq) < 1:
                return True
            body = seq[0]
            for elem in body.find_all("a"):
                href = elem.attrs.get("href", "")
                if len(href) < 1:
                    continue
                if not href.startswith("#"):
                    continue
                href = href[1:]
                db[href] = href
            if len(db) < 1:
                db[""] = "dummy"
            self.bookmarks_anchored = db  # optional
        if s_href not in db:
            return True
        return False


def main(opts: options.Options) -> int:  # {{{1
    logging.basicConfig(level=opts.level_debug)

    common.init(opts.force_offline)
    with open(opts.fname_in, encoding=opts.encoding) as fp:
        data = fp.read()
    prog = HtmlConvertDocx(opts.fname_in)
    prog.on_post_page(data, opts.backend_bs4)
    prog.write_out(opts.fname_out)
    if opts.remove_temporaries:
        common.remove_temporaries()
    return 0


if __name__ == "__main__":  # {{{1
    import sys
    opts = options.parse(sys.argv[1:])
    main(opts)

# vi: ft=python:fdm=marker
