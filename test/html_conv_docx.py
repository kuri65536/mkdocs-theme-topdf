# Copyright (c) 2020, Shimoda <kuri65536 at hotmail dot com>
#
# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at http://mozilla.org/MPL/2.0/.
#
import logging
from logging import (debug as debg, info, warning as warn, )
from lxml import etree  # type: ignore
import sys
from typing import (Dict, Iterable, List, Optional, Text, Tuple, Union, )

from bs4 import BeautifulSoup  # type: ignore
from bs4.element import Tag  # type: ignore

from docx import Document  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore
from docx.shared import Mm  # type: ignore
from docx.table import Table  # type: ignore


if False:
    List


class cfg:
    mode_no_template = True


style_aliases = {
    "Heading 2": "Heading2",
    "Heading 3": "Heading3",
    "Heading 4": "Heading4",
    "Heading 5": "Heading5",
    "Heading 6": "Heading6",

    "List Bullet": "List",
    "List Number": "List",
    "Quote": "Quote",

    "Table Grid": "DefaultTable",
}


def classes_from_prev_sibling(target: Tag) -> Iterable[Text]:  # {{{1
    for elem in target.previous_siblings:
        if elem.name is None:
            continue
        ret: List[Text] = elem.attrs.get("class", [])
        return ret + []
    return []


class HtmlConvertDocx(object):  # {{{1
    def __init__(self) -> None:  # {{{1
        if cfg.mode_no_template:
            doc = Document()
        else:
            doc = Document("template.docx")
        self.output = doc
        info("structure root")
        self.para = doc.add_paragraph('')
        self.header_init()

        if not cfg.mode_no_template:
            # LibreOffice 6.4 can not save Table Styles properly.
            doc.styles.add_style("DefaultTable", WD_STYLE_TYPE.TABLE)

    def header_init(self) -> None:  # {{{1
        # TODO(shimoda): set page border
        sec = self.output.sections[0]
        # moved to template.docx
        if not cfg.mode_no_template:
            return

        # page setting
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = sec.right_margin = sec.bottom_margin = Mm(20)
        sec.top_margin = Mm(20)  # 20 - 1

        # make header to right
        sec.header_distance = Mm(13)  # 20 - 1
        para = self.output.sections[0].header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # page border as rectangle (old word style)
        """<w:pict><v:rect id="shape_0" ID="Shape1" stroked="t"
                style="position:absolute;margin-left:-5.8pt;margin-top:13.05pt;width:493.2pt;height:751.15pt">
                <w10:wrap type="none"/>
                <v:fill o:detectmouseclick="t" on="false"/>
                <v:stroke color="black" weight="12600" joinstyle="round"
                          endcap="flat"/>
           </v:rect></w:pict>
        """
        v = "urn:schemas-microsoft-com:vml"
        w10 = "urn:schemas-microsoft-com:office:word"
        r = para._p.add_r()
        pict = OxmlElement('w:pict')
        r.append(pict)
        rect = etree.Element("{%s}rect" % v)
        pict.append(rect)
        rect.set('id', 'shape_0')  # page_border')
        rect.set('ID', 'Shape1')   # _page_border')
        rect.set('stroked', 't')
        rect.set('style', 'position:absolute;'
                 'margin-left:-5.8pt;margin-top:13.05pt;'
                 'width:493.2pt;height:751.15pt')
        wrap = etree.Element('{%s}wrap' % w10)
        wrap.set('type', "none")
        rect.append(wrap)
        fill = etree.Element('{%s}fill' % v)
        fill.set('on', "false")
        rect.append(fill)
        stroke = etree.Element('{%s}stroke' % v)
        stroke.set('color', "black")
        stroke.set('weight', "12600")
        stroke.set('joinstyle', "round")
        stroke.set('endcap', "flat")
        rect.append(stroke)

    def header_set(self, src: Text) -> None:  # {{{1
        para = self.output.sections[0].header.paragraphs[0]
        para.add_run(src + "( ")
        self.add_field(para, "PAGE")
        para.add_run(" / ")
        self.add_field(para, "NUMPAGES")
        para.add_run(" )")

    def add_field(self, para: Paragraph, instr: Text) -> None:  # {{{1
        r = para.add_run("")._r
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), "begin")
        r.append(fld)

        r = para.add_run("")._r
        cmd = OxmlElement('w:instrText')
        cmd.text = instr
        r.append(cmd)

        r = para.add_run("")._r
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), "separate")
        r.append(fld)

        r = para.add_run("")._r
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), "end")
        r.append(fld)

    def write_out(self) -> None:
        info("structure save")
        fname = 'temp.docx'
        self.output.save(fname)

        doc = Document(fname)
        for para in doc.paragraphs:
            ret = Text(para.text)
            info("after-para: " + (ret[0:10] if len(ret) > 9 else ret))
        for tbl in doc.tables:
            ret = "%d,%d" % (len(tbl.rows), len(tbl.rows[0].cells))
            info("after-tabl: " + (ret))

    def on_post_page(self, output_content: Text, config: Dict[Text, Text],
                     **kwardgs: Text) -> Text:  # {{{1
        soup = BeautifulSoup(output_content, 'html.parser')
        dom = soup.find("body")
        self.extract_para(dom, 0)
        return output_content

    def style(self, name: Text) -> Text:  # {{{1
        if cfg.mode_no_template:
            return name
        if name in style_aliases:
            name = style_aliases[name]
            if name in self.output.styles:
                return name
        warn("can not found style: %s" % name)
        for style in self.output.styles:
            return style.name  # type: ignore
        assert False, "can not found style: %s" % name

    def extract_element(self, elem: Tag) -> Union[None, Text, Tag]:  # {{{1
        if elem.name is None:
            return self.extract_text(elem)

        classes = elem.attrs.get("class", [])
        if "doc-num" in classes:
            self.header_set(elem.string)
            return None
        if "toc" in classes:
            # TODO(shimoda): insert TOC macro.
            return None

        if elem.name == "dl":
            return self.extract_dldtdd(elem)
        elif elem.name == "ul":
            return self.extract_list(elem, False)
        elif elem.name == "ol":
            return self.extract_list(elem, True)
        elif elem.name == "table":
            return self.extract_table(elem)
        elif elem.name == "details":
            return self.extract_details(elem)
        elif elem.name == "svg":
            return self.extract_svg(elem)
        elif elem.name in ("h1", "h2", "h3", "h4", "h5", "h6", ):
            return self.extract_title(elem)
        elif elem.name in ("pre", "code"):
            return self.extract_codeblock(elem)
        # p, article or ...
        return elem

    def extract_text(self, elem: Tag) -> Text:
        ret = Text(elem.string)
        debg(ret.strip())
        return ret

    def extract_table_tree(self, elem: Tag, row: int  # {{{1
                           ) -> Dict[Tuple[int, int], Tag]:
        def num_row(dct: Dict[Tuple[int, int], Tag]) -> int:
            ret = -1
            for (row, col) in dct.keys():
                ret = max(ret, row)
            return ret + 1

        ret: Dict[Tuple[int, int], Tag] = {}
        f_row = False
        for tag in elem.children:
            if tag.name is None:
                continue
            if tag.name == "thead":
                warn("structure: tbl: enter thead")
                ret = self.extract_table_tree(tag, 0)
                continue
            if tag.name == "tbody":
                warn("structure: tbl: enter tbody")
                ret2 = self.extract_table_tree(tag, num_row(ret))
                ret.update(ret2)
                continue
            if tag.name == "tfoot":
                warn("structure: tbl: enter tfoot")
                ret2 = self.extract_table_tree(tag, num_row(ret))
                ret.update(ret2)
                continue
            if tag.name != "tr":
                warn("table tree: %s in table" % tag.name)
                continue
            f_row, col, row = True, 0, (row + 1 if f_row else row)
            for cell in tag.children:
                if cell.name is None:
                    continue
                if cell.name not in ("th", "td", ):
                    warn("table tree: %s in tr" % cell.name)
                    continue
                ret[row, col] = cell
                col += 1
        return ret

    def extract_table(self, elem: Tag) -> Optional[Text]:
        # TODO(shimoda): rowspan, colspan, table in table or styled-cell etc...
        dct = self.extract_table_tree(elem, 0)
        if len(dct) < 1:
            warn("table: did not have any data" + Text(elem.string))
            return None
        n_row = max([tup[0] for tup in dct.keys()]) + 1
        n_col = max([tup[1] for tup in dct.keys()]) + 1

        info("structure: tbl: %s (%d,%d)" % (elem.name, n_row, n_col))
        tbl = self.output.add_table(rows=n_row, cols=n_col)
        for (row, col), cell in sorted(dct.items(), key=lambda x: x[0]):
            src = self.extract_table_cell(cell)
            info("table-%d,%d: %s" % (row, col, src))
            tbl.rows[row].cells[col].text = src
        return None

    def extract_dldtdd(self, elem: Tag) -> Optional[Text]:
        classes = classes_from_prev_sibling(elem)
        if "before-dl-table" not in classes:
            import pdb
            pdb.set_trace()

        n_row = n_col = i = 0
        for tag in elem.children:
            if tag.name not in ("dt", "dd"):
                continue
            if tag.name == "dt":
                n_row, i = n_row + 1, 1
                continue
            i += 1
            n_col = max(i, n_col)

        info("structure: tbl: %s (%d,%d)" % (elem.name, n_row, n_col))
        tbl = self.output.add_table(rows=n_row, cols=n_col)
        tbl.autofit = False
        tbl.style = self.style("Table Grid")
        n_row, i = -1, 0
        for tag in elem.children:
            if tag.name not in ("dt", "dd"):
                continue
            src = self.extract_table_cell(tag)
            if tag.name == "dt":
                n_row, i = n_row + 1, 0
                debg("dt-%d,%d: %s" % (n_row, i, src))
                tbl.rows[n_row].cells[i].text = src
                continue
            i += 1
            debg("dd-%d,%d: %s" % (n_row, i, src))
            tbl.rows[n_row].cells[i].text = src

        if self.style_table_stamps(tbl, classes):
            return None
        self.style_table_width_from(tbl, classes)
        return None

    def extract_list(self, elem: Tag, f_number: bool) -> Optional[Text]:
        style = "List Number" if f_number else "List Bullet"
        style = self.style(style)
        for tag in elem.children:
            if tag.name != "li":
                continue
            ret = Text(tag.text)
            info("structure: li : " + ret.splitlines()[0])
            self.output.add_paragraph(ret, style=style)
        return None

    def extract_details(self, elem: Tag) -> Optional[Text]:
        debg("structure: details: not implemented, skipped...")
        return None

    def extract_svg(self, elem: Tag) -> Optional[Text]:
        info("structure: svg: not implemented, skipped...")
        # TODO(shimoda): impl
        # 1. write_out_svg
        # fp = open("temp.svg", "w")
        # fp.close()
        # self.output.add_picture("temp.svg")
        return None

    def extract_title(self, elem: Tag) -> Optional[Text]:
        level = int(elem.name.lstrip("h"))
        ret = Text(elem.text)
        info("structure: hed: " + ret)
        style = self.style("Heading " + Text(level))
        self.output.add_paragraph(ret, style=style)
        return None

    def extract_codeblock(self, elem: Tag) -> Optional[Text]:
        ret = Text(elem.string)
        info("structure: pre: " + ret.splitlines()[0])
        style = style_aliases["Quote"]
        para = self.output.add_paragraph(ret, style=style)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        for val in ["left", "right", "top", "bottom"]:
            b = OxmlElement('w:' + val)
            b.set(qn('w:val'), 'thinThickLargeGap')
            b.set(qn('w:sz'), '2')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), '000000')
            pBdr.append(b)
        return None

    def extract_para(self, node: Tag, level: int) -> Optional[Text]:
        debg("enter para...: %d-%s" % (level, node.name))
        # para = self.para
        for elem in node.children:
            ret = self.extract_element(elem)
            if isinstance(ret, Text):
                pass
                # para.text += ret
            elif ret is not None:
                self.extract_para(elem, level + 1)
            # else:
            #     para = self.output.add_paragraph('')
        return None

    def extract_table_cell(self, node: Tag) -> Text:
        ret = ""
        for elem in node:
            if elem.name is None:
                ret += elem.string
            elif elem.name == "br":
                ret += "\n"
            else:
                debg("can't extract %s in a table-cell" % elem.name)
        return ret

    def style_table_stamps(self, tbl: Table,  # {{{1
                           classes: Iterable[Text]
                           ) -> bool:
        if "table-3stamps" not in classes:
            return False
        for col, wid in zip(tbl.columns,
                            [Mm(106), Mm(20), Mm(20), Mm(20)]):
            col.width = wid
        tbl.rows[0].height = Mm(20)
        # TODO(shimoda): col1 -> smaller at 1st line and large after
        # TODO(shimoda): col2-4 -> smaller font,
        return True

    def style_table_width_from(self, tbl: Table,  # {{{1
                               classes: Iterable[Text]
                               ) -> bool:
        ret = {"table2-8": [Mm(32), Mm(138)],
               "table3-7": [Mm(48), Mm(112)],
               "table4-6": [Mm(64), Mm(96)],
               "table5-5": [Mm(80), Mm(80)],
               "table2-2-6": [Mm(32), Mm(32), Mm(96)],
               "table2-3-5": [Mm(32), Mm(48), Mm(80)],
               "table2-4-4": [Mm(32), Mm(64), Mm(64)],
               "table2-5-3": [Mm(32), Mm(80), Mm(48)],
               "table2-6-2": [Mm(32), Mm(96), Mm(32)],
               "table3-3-3": [Mm(53), Mm(53), Mm(53)],
               "table2-2-2-4": [Mm(32), Mm(32), Mm(32), Mm(64)],
               }

        widths: List[Mm] = []
        for class_ in classes:
            if class_ in ret:
                widths = ret[class_]
                break
        else:
            warn("width did not specified by class")
            return False
        for col, wid in zip(tbl.columns, widths):
            col.width = wid
        return True


def main() -> int:  # {{{1
    logging.basicConfig(level=logging.INFO)

    opts = sys.argv[1:]
    data = open(opts[0]).read()
    prog = HtmlConvertDocx()
    prog.on_post_page(data, {})
    prog.write_out()
    return 0


if __name__ == "__main__":  # {{{1
    main()

# vi: ft=python
